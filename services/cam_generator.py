"""
CAM Generator Service — Enhanced
------------------------------------
Generates a professional Credit Appraisal Memo as:
  - Word (.docx) via python-docx
  - PDF (.pdf)  via reportlab

Sections:
  1. Executive Summary
  2. Company & Promoter Profile
  3. Financial Analysis with Indian Ratios
  4. Five Cs Assessment with detailed breakdown
  5. GST & Bank Statement Analysis
  6. Risk Flags & Early Warning Signals
  7. Regulatory & Compliance Assessment
  8. Credit Officer Observations (if available)
  9. Recommendation & Loan Terms
  10. Annexure — Scoring Methodology
"""

import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table,
    TableStyle, HRFlowable, PageBreak
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT


def _load_json(path: str) -> dict:
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {}


# ─────────────────────────────────────────────
# INDIAN FINANCIAL RATIO CALCULATOR
# ─────────────────────────────────────────────

def compute_ratios(financial_data: dict, score_data: dict) -> dict:
    """Compute Indian-context financial ratios from available data."""
    revenue     = financial_data.get("revenue", 0)
    ebitda      = financial_data.get("ebitda", 0)
    net_profit  = financial_data.get("net_profit", 0)
    total_debt  = financial_data.get("total_debt", 0)
    net_worth   = financial_data.get("net_worth", 0)
    total_assets = financial_data.get("total_assets", net_worth + total_debt) or 1
    current_assets = financial_data.get("current_assets", 0)
    current_liabilities = financial_data.get("current_liabilities", 0)
    receivables = financial_data.get("receivables", 0)
    inventory   = financial_data.get("inventory", 0)
    cogs        = financial_data.get("cogs", revenue * 0.65 if revenue else 0)
    interest    = financial_data.get("interest_expense", total_debt * 0.10 if total_debt else 0)

    ratios = {}

    # Profitability
    ratios["ebitda_margin"]     = round((ebitda / revenue) * 100, 2)      if revenue > 0 and ebitda else None
    ratios["net_profit_margin"] = round((net_profit / revenue) * 100, 2)  if revenue > 0 and net_profit else None

    # Return Ratios
    ratios["roe"]  = round((net_profit / net_worth) * 100, 2)            if net_worth > 0 and net_profit else None
    ratios["roa"]  = round((net_profit / total_assets) * 100, 2)         if total_assets > 0 and net_profit else None
    ratios["roce"] = round((ebitda / (net_worth + total_debt)) * 100, 2) if (net_worth + total_debt) > 0 and ebitda else None

    # Leverage
    ratios["debt_equity"]    = round(total_debt / net_worth, 2)     if net_worth > 0 else None
    ratios["debt_to_assets"] = round(total_debt / total_assets, 2)  if total_assets > 0 else None

    # Liquidity
    ratios["current_ratio"] = round(current_assets / current_liabilities, 2)                  if current_liabilities > 0 else None
    ratios["quick_ratio"]   = round((current_assets - inventory) / current_liabilities, 2)    if current_liabilities > 0 and current_assets > inventory else None

    # Coverage
    ratios["dscr"]              = score_data.get("dscr", 0)
    ratios["interest_coverage"] = round(ebitda / interest, 2)  if interest > 0 and ebitda else None

    # Efficiency
    ratios["dso"]            = round((receivables / revenue) * 365, 0) if revenue > 0 and receivables else None
    ratios["asset_turnover"] = round(revenue / total_assets, 2)        if total_assets > 0 and revenue else None
    ratios["inventory_days"] = round((inventory / cogs) * 365, 0)      if cogs > 0 and inventory else None

    return ratios


def ratio_benchmark(name: str, value) -> dict:
    """Return benchmark and status for a given ratio."""
    if value is None:
        return {"benchmark": "N/A", "status": "No Data", "color": "grey"}

    benchmarks = {
        "current_ratio":     {"good": 1.33, "fair": 1.0,  "label": "> 1.33x"},
        "quick_ratio":       {"good": 1.0,  "fair": 0.7,  "label": "> 1.0x"},
        "debt_equity":       {"good": 1.5,  "fair": 2.5,  "label": "< 1.5x",     "inverted": True},
        "dscr":              {"good": 1.5,  "fair": 1.25, "label": "> 1.5x"},
        "interest_coverage": {"good": 3.0,  "fair": 2.0,  "label": "> 3.0x"},
        "roe":               {"good": 15,   "fair": 10,   "label": "> 15%"},
        "roa":               {"good": 5,    "fair": 3,    "label": "> 5%"},
        "roce":              {"good": 15,   "fair": 10,   "label": "> 15%"},
        "ebitda_margin":     {"good": 15,   "fair": 10,   "label": "> 15%"},
        "net_profit_margin": {"good": 10,   "fair": 5,    "label": "> 10%"},
        "dso":               {"good": 60,   "fair": 90,   "label": "< 60 days",   "inverted": True},
        "asset_turnover":    {"good": 1.5,  "fair": 1.0,  "label": "> 1.5x"},
        "inventory_days":    {"good": 60,   "fair": 90,   "label": "< 60 days",   "inverted": True},
        "debt_to_assets":    {"good": 0.4,  "fair": 0.6,  "label": "< 0.4x",      "inverted": True},
    }

    b = benchmarks.get(name, {"good": 50, "fair": 25, "label": "N/A"})
    inverted = b.get("inverted", False)

    if inverted:
        if value <= b["good"]:   return {"benchmark": b["label"], "status": "Strong",   "color": "green"}
        elif value <= b["fair"]: return {"benchmark": b["label"], "status": "Adequate",  "color": "amber"}
        else:                    return {"benchmark": b["label"], "status": "Weak",      "color": "red"}
    else:
        if value >= b["good"]:   return {"benchmark": b["label"], "status": "Strong",   "color": "green"}
        elif value >= b["fair"]: return {"benchmark": b["label"], "status": "Adequate",  "color": "amber"}
        else:                    return {"benchmark": b["label"], "status": "Weak",      "color": "red"}


# ─────────────────────────────────────────────
# WORD DOCUMENT GENERATOR
# ─────────────────────────────────────────────

def generate_cam_docx(company_id: str, score_data: dict, research_data: dict) -> str:
    """Generate a comprehensive Word (.docx) CAM."""
    os.makedirs(f"uploads/{company_id}", exist_ok=True)
    output_path = f"uploads/{company_id}/cam_report.docx"

    parse_data    = _load_json(f"uploads/{company_id}/parse_results.json")
    gst_flags     = _load_json(f"uploads/{company_id}/gst_flags.json")
    officer_notes = _load_json(f"uploads/{company_id}/officer_notes.json")

    financial_data = {}
    for d in parse_data.get("parsed_docs", []):
        financial_data.update(d.get("financials", {}))

    ratios = compute_ratios(financial_data, score_data)

    doc = Document()
    _set_margins(doc)

    # Header
    _heading_center(doc, "INTELLI-CREDIT FINANCIAL SERVICES", 16, RGBColor(0x1a, 0x4a, 0x7a))
    _heading_center(doc, "AI-Powered Corporate Credit Appraisal Engine", 9, RGBColor(0x99, 0x99, 0x99))
    _heading_center(doc, "CREDIT APPRAISAL MEMORANDUM (CAM)", 14, RGBColor(0, 0, 0))
    _add_meta_table(doc, score_data, research_data)

    # Sections
    _section(doc, "1. EXECUTIVE SUMMARY")
    _exec_summary(doc, score_data, research_data, ratios)

    _section(doc, "2. COMPANY & PROMOTER PROFILE")
    _company_profile(doc, research_data, parse_data)

    _section(doc, "3. FINANCIAL ANALYSIS")
    _financial_analysis(doc, score_data, financial_data, ratios)

    _section(doc, "4. FIVE Cs CREDIT ASSESSMENT")
    _five_cs_docx(doc, score_data)

    _section(doc, "5. GST & BANK STATEMENT ANALYSIS")
    _gst_analysis(doc, gst_flags)

    _section(doc, "6. RISK FLAGS & EARLY WARNING SIGNALS")
    _risk_flags_docx(doc, research_data, score_data)

    _section(doc, "7. REGULATORY & COMPLIANCE ASSESSMENT")
    _regulatory_docx(doc, research_data)

    sec = 8
    if officer_notes:
        _section(doc, f"{sec}. CREDIT OFFICER OBSERVATIONS")
        _officer_notes_docx(doc, officer_notes)
        sec += 1

    _section(doc, f"{sec}. RECOMMENDATION & LOAN TERMS")
    _recommendation_docx(doc, score_data, ratios)

    _section(doc, f"{sec+1}. ANNEXURE — SCORING METHODOLOGY")
    _methodology_docx(doc)

    doc.save(output_path)
    return output_path


# ── Word helpers ──

def _set_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.8)
        s.left_margin = s.right_margin = Inches(1.0)

def _heading_center(doc, text, size, color):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = size > 10
    r.font.size = Pt(size)
    r.font.color.rgb = color

def _section(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.underline = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1a, 0x4a, 0x7a)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def _add_meta_table(doc, score_data, research_data):
    t = doc.add_table(rows=5, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cid = score_data.get("company_id", "N/A")
    d = score_data.get("decision", "PENDING")
    di = "✅ APPROVE" if d == "APPROVE" else ("❌ REJECT" if d == "REJECT" else "⚠️ REFER")
    ra = score_data.get("recommended_amount")
    rows = [
        ("Company ID", cid, "Date", datetime.now().strftime("%d %B %Y")),
        ("Decision", di, "Credit Score", f"{score_data.get('overall_score',0)}/100"),
        ("Interest Rate", score_data.get("interest_rate","N/A"), "Rec. Amount", f"₹{ra:,.0f}" if ra else "N/A"),
        ("Research Risk", research_data.get("overall_research_risk","N/A"), "DSCR", str(score_data.get("dscr","N/A"))),
        ("Prepared By", "Intelli-Credit AI", "Reviewed By", "Credit Committee"),
    ]
    for i,(l1,v1,l2,v2) in enumerate(rows):
        t.rows[i].cells[0].text = l1
        t.rows[i].cells[1].text = str(v1)
        t.rows[i].cells[2].text = l2
        t.rows[i].cells[3].text = str(v2)
        for c in [0,2]:
            t.rows[i].cells[c].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

def _exec_summary(doc, score_data, research_data, ratios):
    d = score_data.get("decision","PENDING")
    s = score_data.get("overall_score",0)
    ra = score_data.get("recommended_amount",0)
    rate = score_data.get("interest_rate","N/A")
    summary = research_data.get("research_summary","Research summary not available.")
    reason = score_data.get("decision_reason","")
    txt = (
        f"This Credit Appraisal Memo presents findings of an AI-powered due diligence exercise "
        f"encompassing financial statements, GST filings, bank statements, and web-scale secondary "
        f"research (promoter background, litigation, sector outlook, MCA/RBI compliance).\n\n"
        f"Overall Credit Score: {s}/100\nDecision: {d}\nRationale: {reason}\n\n"
        f"Key Financial Indicators:\n"
    )
    for k,l in [("dscr","DSCR"),("current_ratio","Current Ratio"),("debt_equity","Debt/Equity"),("roe","ROE"),("ebitda_margin","EBITDA Margin")]:
        v = ratios.get(k)
        if v is not None:
            unit = "%" if k in ("roe","ebitda_margin","net_profit_margin","roa","roce") else "x"
            txt += f"  • {l}: {v}{unit}\n"
    txt += f"\nResearch Summary: {summary}\n"
    if d == "APPROVE" and ra:
        txt += f"\nRecommendation: Approve facility of ₹{ra:,.0f} at {rate}."
    elif d == "REFER":
        txt += "\nRecommendation: Refer to credit committee for further review."
    else:
        txt += f"\nRecommendation: {reason}"
    doc.add_paragraph(txt)

def _company_profile(doc, research_data, parse_data):
    doc.add_paragraph(f"2.1 Promoter Background:\n{research_data.get('promoter_background','Not available.')}")
    doc.add_paragraph(f"\n2.2 MCA/ROC Compliance:\n{research_data.get('mca_compliance','Not available.')}")
    doc.add_paragraph(f"\n2.3 Document Insights:\n{research_data.get('document_insights','Not available.')}")
    parsed = parse_data.get("parsed_docs",[])
    if parsed:
        doc.add_paragraph("\n2.4 Documents Analysed:")
        for pd_doc in parsed:
            doc.add_paragraph(
                f"  • {pd_doc.get('doc_type','Unknown')}: {pd_doc.get('file','N/A')} "
                f"({pd_doc.get('pages',0)} pages)", style="List Bullet")

def _financial_analysis(doc, score_data, financial_data, ratios):
    doc.add_paragraph("3.1 Key Financial Figures:")
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
    for i,h in enumerate(["Metric","Value (₹)","Source"]):
        t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for label,key,src in [("Revenue","revenue","Annual Report / GST"),("EBITDA","ebitda","Annual Report"),
                           ("Net Profit","net_profit","P&L"),("Total Debt","total_debt","Balance Sheet"),
                           ("Net Worth","net_worth","Balance Sheet")]:
        v = financial_data.get(key,0)
        row = t.add_row(); row.cells[0].text = label; row.cells[1].text = f"₹{v:,.0f}" if v else "N/A"; row.cells[2].text = src
    doc.add_paragraph()

    doc.add_paragraph("3.2 Financial Ratio Analysis (Indian Benchmarks):")
    t2 = doc.add_table(rows=1, cols=5); t2.style = "Table Grid"
    for i,h in enumerate(["Ratio","Value","Benchmark","Status","Category"]):
        t2.rows[0].cells[i].text = h; t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for label,key,cat in [("Current Ratio","current_ratio","Liquidity"),("Quick Ratio","quick_ratio","Liquidity"),
                           ("Debt/Equity","debt_equity","Leverage"),("Debt-to-Assets","debt_to_assets","Leverage"),
                           ("DSCR","dscr","Coverage"),("Interest Coverage","interest_coverage","Coverage"),
                           ("EBITDA Margin","ebitda_margin","Profitability"),("Net Margin","net_profit_margin","Profitability"),
                           ("ROE","roe","Returns"),("ROA","roa","Returns"),("ROCE","roce","Returns"),
                           ("DSO (Days)","dso","Efficiency"),("Asset Turnover","asset_turnover","Efficiency"),
                           ("Inventory Days","inventory_days","Efficiency")]:
        v = ratios.get(key)
        bm = ratio_benchmark(key, v)
        row = t2.add_row(); row.cells[0].text = label; row.cells[1].text = str(v) if v is not None else "N/A"
        row.cells[2].text = bm["benchmark"]; row.cells[3].text = bm["status"]; row.cells[4].text = cat
    doc.add_paragraph()

    dscr = ratios.get("dscr",0)
    doc.add_paragraph("3.3 DSCR Commentary:")
    if dscr >= 1.5:
        doc.add_paragraph(f"DSCR of {dscr}x indicates strong debt servicing capacity with adequate safety margin.")
    elif dscr >= 1.25:
        doc.add_paragraph(f"DSCR of {dscr}x is adequate but below ideal 1.5x threshold. Monitor closely.")
    elif dscr > 0:
        doc.add_paragraph(f"⚠️ DSCR of {dscr}x below acceptable threshold. Enhanced monitoring required.")
    else:
        doc.add_paragraph("DSCR could not be calculated — missing financial data.")

def _five_cs_docx(doc, score_data):
    cs = [
        ("Character",  score_data.get("character_score",0),  "25%", "Promoter integrity, litigation, MCA, credit track record"),
        ("Capacity",   score_data.get("capacity_score",0),   "30%", "DSCR, revenue, EBITDA, GST compliance, cash flow"),
        ("Capital",    score_data.get("capital_score",0),    "20%", "Net worth, leverage, debt-equity, capital cushion"),
        ("Collateral", score_data.get("collateral_score",0), "15%", "Security coverage, asset quality, enforceability"),
        ("Conditions", score_data.get("conditions_score",0), "10%", "Sector outlook, RBI regs, macro conditions"),
    ]
    t = doc.add_table(rows=len(cs)+2, cols=4); t.style = "Table Grid"
    for i,h in enumerate(["Five C","Weight","Score (/100)","Scope"]):
        t.rows[0].cells[i].text = h; t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i,(n,sc,w,d) in enumerate(cs,1):
        icon = "✅" if sc >= 70 else "⚠️" if sc >= 50 else "❌"
        r = t.rows[i]; r.cells[0].text = n; r.cells[1].text = w; r.cells[2].text = f"{sc}/100 {icon}"; r.cells[3].text = d
    ov = score_data.get("overall_score",0)
    icon = "✅" if ov >= 70 else "⚠️" if ov >= 50 else "❌"
    lr = t.rows[len(cs)+1]
    lr.cells[0].text = "OVERALL"; lr.cells[1].text = "100%"; lr.cells[2].text = f"{ov}/100 {icon}"; lr.cells[3].text = "Weighted aggregate"
    for c in range(4): lr.cells[c].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    ded = score_data.get("deductions",{})
    if ded:
        doc.add_paragraph("Score Deduction Audit Trail:")
        for cat,items in ded.items():
            if items:
                doc.add_paragraph(f"\n[{cat.upper()}]:")
                for item in items:
                    doc.add_paragraph(f"  • {item}", style="List Bullet")

def _gst_analysis(doc, gst_flags):
    if gst_flags:
        fr = gst_flags.get("fraud_risk","LOW"); rr = gst_flags.get("revenue_inflation_risk","LOW")
        ct = gst_flags.get("circular_trading_detected",False)
        t = doc.add_table(rows=4, cols=2); t.style = "Table Grid"
        for i,(l,v) in enumerate([
            ("ITC Fraud Risk (GSTR-2B vs 3B)", f"{'❌' if fr=='HIGH' else '⚠️' if fr=='MEDIUM' else '✅'} {fr}"),
            ("Revenue Inflation Risk", f"{'❌' if rr=='HIGH' else '⚠️' if rr=='MEDIUM' else '✅'} {rr}"),
            ("Circular Trading", f"{'❌ DETECTED' if ct else '✅ Not Detected'}"),
            ("Summary", gst_flags.get("summary","N/A")),
        ]):
            t.rows[i].cells[0].text = l; t.rows[i].cells[1].text = v
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    else:
        doc.add_paragraph("GST validation data not available. GSTR-2B vs 3B cross-check pending.")
    doc.add_paragraph(
        "\nNote: GSTR-2B (auto-populated ITC) is cross-checked against GSTR-3B (self-declared ITC). "
        "Overclaims indicate potential ITC fraud. Bank credits vs GST turnover flags revenue inflation."
    )

def _risk_flags_docx(doc, research_data, score_data):
    signals = research_data.get("early_warning_signals",[])
    lf = research_data.get("litigation_flags",[])
    pos = research_data.get("positive_factors",[])
    if signals:
        doc.add_paragraph("⚠️ Early Warning Signals:")
        for s in signals: doc.add_paragraph(f"  • {s}", style="List Bullet")
    if lf:
        doc.add_paragraph("⚖️ Litigation & Legal Flags:")
        for f in lf: doc.add_paragraph(f"  • {f}", style="List Bullet")
    if pos:
        doc.add_paragraph("✅ Positive Factors:")
        for p in pos: doc.add_paragraph(f"  • {p}", style="List Bullet")
    if not signals and not lf:
        doc.add_paragraph("No critical risk flags identified.")
    r = research_data.get("overall_research_risk","MEDIUM")
    doc.add_paragraph(f"\nRisk Summary — Overall: {r} | Warnings: {len(signals)} | Litigation: {len(lf)} | Positives: {len(pos)}")

def _regulatory_docx(doc, research_data):
    doc.add_paragraph(f"7.1 RBI & Regulatory Risks:\n{research_data.get('regulatory_risks','No specific risks identified.')}")
    doc.add_paragraph(f"\n7.2 Sector Outlook:\n{research_data.get('sector_outlook','Not available.')}")
    doc.add_paragraph("\n7.3 Key Indian Regulatory Framework:")
    for r in [
        "RBI IRAC Norms — Income Recognition, Asset Classification & Provisioning",
        "RBI Large Exposure Framework (LEF) — Borrower group exposure limits",
        "SEBI corporate governance & disclosure norms (listed entities)",
        "MCA compliance — Annual filing, Director KYC, charge registration",
        "GST — Continuous GSTR-2A/2B vs GSTR-3B reconciliation for ITC",
        "FEMA regulations for foreign exchange transactions",
        "IBC/NCLT proceedings — CIRP check for insolvency history",
        "CIBIL Commercial credit report — Payment history & defaults",
    ]:
        doc.add_paragraph(f"  • {r}", style="List Bullet")

def _officer_notes_docx(doc, notes):
    doc.add_paragraph(f"Site Visit Date: {notes.get('site_visit_date','N/A')}")
    doc.add_paragraph(f"Observations:\n{notes.get('notes','No notes.')}")
    doc.add_paragraph("Note: These qualitative inputs have been factored into AI scoring via adjustment factors.")

def _recommendation_docx(doc, score_data, ratios):
    d = score_data.get("decision","PENDING"); s = score_data.get("overall_score",0)
    reason = score_data.get("decision_reason",""); ra = score_data.get("recommended_amount",0)
    rate = score_data.get("interest_rate","N/A")
    doc.add_paragraph(f"Decision: {d}\nOverall Credit Score: {s}/100\nRationale: {reason}")
    if d == "APPROVE" and ra:
        doc.add_paragraph("\nProposed Loan Terms:")
        t = doc.add_table(rows=7, cols=2); t.style = "Table Grid"
        for i,(l,v) in enumerate([
            ("Facility Amount", f"₹{ra:,.0f}"), ("Interest Rate", rate),
            ("Facility Type", "Working Capital / Term Loan"), ("Tenor", "12 Months (Renewable)"),
            ("Primary Security", "Hypothecation of current assets + Fixed assets"),
            ("Collateral", "Property / FD per bank norms"),
            ("Conditions", "Legal docs, security creation, CC approval"),
        ]):
            t.rows[i].cells[0].text = l; t.rows[i].cells[1].text = v
            t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    elif d == "REFER":
        doc.add_paragraph("Referred to Credit Committee for detailed review.")
    else:
        doc.add_paragraph("NOT recommended for approval. See risk flags above.")
    doc.add_paragraph(f"\nFull Explanation:\n{score_data.get('explanation','')}")

def _methodology_docx(doc):
    doc.add_paragraph(
        "The Intelli-Credit model uses transparent, rule-based Five Cs scoring. "
        "Each category starts at 100 with deductions for identified risks.\n"
    )
    t = doc.add_table(rows=6, cols=3); t.style = "Table Grid"
    for i,(c1,c2,c3) in enumerate([
        ("Category","Weight","Key Factors"),
        ("Character","25%","Promoter, NCLT/IBC, MCA, research risk"),
        ("Capacity","30%","DSCR, revenue, EBITDA, GST fraud, circular trading"),
        ("Capital","20%","D/E ratio, net worth, leverage"),
        ("Collateral","15%","Security coverage vs facility amount"),
        ("Conditions","10%","Sector outlook, RBI regulatory risk"),
    ]):
        t.rows[i].cells[0].text = c1; t.rows[i].cells[1].text = c2; t.rows[i].cells[2].text = c3
        if i == 0:
            for c in range(3): t.rows[i].cells[c].paragraphs[0].runs[0].bold = True
    doc.add_paragraph(
        "\nDecision Rules: APPROVE (≥70) | REFER (50–70) | REJECT (<50 or hard-reject)\n"
        "Hard Rejects: NCLT/IBC, Circular trading, Character <30, Capacity <25\n"
        "Rate: ≥80 → MCLR+1% | 65–80 → MCLR+2% | 50–65 → MCLR+3.5%"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Generated by Intelli-Credit AI | {datetime.now().strftime('%d %B %Y, %I:%M %p')} | CONFIDENTIAL")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x99,0x99,0x99)


# ─────────────────────────────────────────────
# PDF GENERATOR (ReportLab) — Enhanced
# ─────────────────────────────────────────────

def generate_cam_pdf(company_id: str, score_data: dict, research_data: dict) -> str:
    """Generate comprehensive PDF CAM."""
    os.makedirs(f"uploads/{company_id}", exist_ok=True)
    output_path = f"uploads/{company_id}/cam_report.pdf"

    parse_data = _load_json(f"uploads/{company_id}/parse_results.json")
    gst_flags  = _load_json(f"uploads/{company_id}/gst_flags.json")
    financial_data = {}
    for d in parse_data.get("parsed_docs", []):
        financial_data.update(d.get("financials", {}))
    ratios = compute_ratios(financial_data, score_data)

    pdf = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=45, leftMargin=45, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    navy = colors.HexColor("#1a4a7a"); red = colors.HexColor("#c0392b")
    green = colors.HexColor("#1a6b3c"); amber_c = colors.HexColor("#b5500a")

    title_s  = ParagraphStyle("T", parent=styles["Title"], textColor=navy, fontSize=18, spaceAfter=4)
    sub_s    = ParagraphStyle("Sub", parent=styles["Normal"], textColor=colors.grey, fontSize=9, alignment=TA_CENTER, spaceAfter=8)
    h1_s     = ParagraphStyle("H1", parent=styles["Heading1"], textColor=navy, fontSize=12, spaceAfter=4, spaceBefore=14)
    h2_s     = ParagraphStyle("H2", parent=styles["Heading2"], textColor=navy, fontSize=10, spaceAfter=3, spaceBefore=8)
    body_s   = ParagraphStyle("B", parent=styles["Normal"], fontSize=9, spaceAfter=5, leading=13)
    small_s  = ParagraphStyle("Sm", parent=styles["Normal"], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
    bullet_s = ParagraphStyle("Bu", parent=styles["Normal"], fontSize=9, spaceAfter=3, leading=12, leftIndent=20, bulletIndent=10)

    el = []
    d = score_data.get("decision","PENDING")
    ov = score_data.get("overall_score",0)
    cid = score_data.get("company_id", company_id)
    dt = datetime.now().strftime("%d %B %Y")
    ra = score_data.get("recommended_amount",0)
    rate = score_data.get("interest_rate","N/A")
    reason = score_data.get("decision_reason","")
    dc = green if d == "APPROVE" else (red if d == "REJECT" else amber_c)

    # Header
    el.append(Paragraph("INTELLI-CREDIT FINANCIAL SERVICES", title_s))
    el.append(Paragraph("AI-Powered Corporate Credit Appraisal Engine", sub_s))
    el.append(Paragraph("Credit Appraisal Memorandum (CAM)", styles["Heading2"]))
    el.append(HRFlowable(width="100%", thickness=2, color=navy))
    el.append(Spacer(1, 8))

    # Meta
    mt = Table([
        ["Company ID", cid, "Date", dt],
        ["Credit Score", f"{ov}/100", "Decision", d],
        ["Interest Rate", rate, "Rec. Amount", f"₹{ra:,.0f}" if ra else "N/A"],
        ["Research Risk", research_data.get("overall_research_risk","N/A"), "DSCR", str(score_data.get("dscr","N/A"))],
    ], colWidths=[1.3*inch, 2*inch, 1.3*inch, 2*inch])
    mt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#f0f4f8")),
        ("BACKGROUND",(2,0),(2,-1),colors.HexColor("#f0f4f8")),
        ("TEXTCOLOR",(0,0),(0,-1),navy),("TEXTCOLOR",(2,0),(2,-1),navy),
        ("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),("FONTNAME",(2,0),(2,-1),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),9),("GRID",(0,0),(-1,-1),0.5,colors.lightgrey),("PADDING",(0,0),(-1,-1),5),
        ("TEXTCOLOR",(3,1),(3,1),dc),("FONTNAME",(3,1),(3,1),"Helvetica-Bold"),
    ]))
    el.append(mt); el.append(Spacer(1,12))

    # S1: Exec Summary
    el.append(Paragraph("1. EXECUTIVE SUMMARY", h1_s))
    summary = research_data.get("research_summary","N/A")
    el.append(Paragraph(f"Score: <b>{ov}/100</b>. Decision: <b>{d}</b>. {reason}", body_s))
    el.append(Paragraph(f"Research: {summary}", body_s))
    el.append(Spacer(1,6))

    # S2: Five Cs
    el.append(Paragraph("2. FIVE Cs CREDIT SCORECARD", h1_s))
    cs = [["Five C","Weight","Score","Rating","Key Factor"],
          ["Character","25%",f"{score_data.get('character_score',0)}/100",_rating_badge(score_data.get('character_score',0)),"Promoter, Litigation"],
          ["Capacity","30%",f"{score_data.get('capacity_score',0)}/100",_rating_badge(score_data.get('capacity_score',0)),"DSCR, Revenue, GST"],
          ["Capital","20%",f"{score_data.get('capital_score',0)}/100",_rating_badge(score_data.get('capital_score',0)),"D/E, Net Worth"],
          ["Collateral","15%",f"{score_data.get('collateral_score',0)}/100",_rating_badge(score_data.get('collateral_score',0)),"Coverage Ratio"],
          ["Conditions","10%",f"{score_data.get('conditions_score',0)}/100",_rating_badge(score_data.get('conditions_score',0)),"Sector, RBI"],
          ["OVERALL","100%",f"{ov}/100",_rating_badge(ov),"Weighted Aggregate"]]
    ct = Table(cs, colWidths=[1.1*inch,.7*inch,.9*inch,1.1*inch,1.8*inch])
    ct.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),navy),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("BACKGROUND",(0,-1),(-1,-1),colors.HexColor("#eaf4ff")),("FONTNAME",(0,-1),(-1,-1),"Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,1),(-1,-2),[colors.white,colors.HexColor("#f8f8f8")]),
        ("GRID",(0,0),(-1,-1),0.5,colors.lightgrey),("PADDING",(0,0),(-1,-1),6),("FONTSIZE",(0,0),(-1,-1),9),
    ]))
    el.append(ct); el.append(Spacer(1,10))

    # S3: Financial Ratios
    el.append(Paragraph("3. KEY FINANCIAL RATIOS (Indian Benchmarks)", h1_s))
    rd = [["Ratio","Value","Benchmark","Status"]]
    for label,key in [("Current Ratio","current_ratio"),("Debt/Equity","debt_equity"),("DSCR","dscr"),
                       ("Interest Coverage","interest_coverage"),("EBITDA Margin","ebitda_margin"),
                       ("Net Margin","net_profit_margin"),("ROE","roe"),("ROCE","roce"),
                       ("DSO (Days)","dso"),("Asset Turnover","asset_turnover")]:
        v = ratios.get(key); bm = ratio_benchmark(key,v)
        rd.append([label, str(v) if v is not None else "N/A", bm["benchmark"], bm["status"]])
    rt = Table(rd, colWidths=[1.5*inch,1*inch,1.5*inch,1.5*inch])
    rt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),navy),("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#f8f8f8")]),
        ("GRID",(0,0),(-1,-1),0.5,colors.lightgrey),("PADDING",(0,0),(-1,-1),5),("FONTSIZE",(0,0),(-1,-1),9),
    ]))
    el.append(rt); el.append(Spacer(1,10))

    # S4: Risk Flags
    el.append(Paragraph("4. RISK FLAGS & EARLY WARNING SIGNALS", h1_s))
    signals = research_data.get("early_warning_signals",[])
    lfl = research_data.get("litigation_flags",[])
    pos = research_data.get("positive_factors",[])
    if signals:
        el.append(Paragraph("<b>Early Warning Signals:</b>", body_s))
        for s in signals: el.append(Paragraph(f"&bull; {s}", bullet_s))
    if lfl:
        el.append(Paragraph("<b>Litigation Flags:</b>", body_s))
        for f in lfl: el.append(Paragraph(f"&bull; {f}", bullet_s))
    if pos:
        el.append(Spacer(1,4)); el.append(Paragraph("<b>Positive Factors:</b>", body_s))
        for p in pos: el.append(Paragraph(f"&bull; {p}", bullet_s))
    if not signals and not lfl:
        el.append(Paragraph("No critical risk flags identified.", body_s))
    el.append(Spacer(1,10))

    # S5: Regulatory
    el.append(Paragraph("5. REGULATORY & COMPLIANCE", h1_s))
    el.append(Paragraph(research_data.get("regulatory_risks","No specific risks."), body_s))
    el.append(Paragraph(f"Sector: {research_data.get('sector_outlook','N/A')[:300]}", body_s))
    el.append(Spacer(1,8))

    # S6: Recommendation
    el.append(Paragraph("6. RECOMMENDATION & LOAN TERMS", h1_s))
    rc = "#1a6b3c" if d == "APPROVE" else ("#c0392b" if d == "REJECT" else "#b5500a")
    el.append(Paragraph(f'<font color="{rc}"><b>DECISION: {d}</b></font>', body_s))
    el.append(Paragraph(f"Reason: {reason}", body_s))
    if d == "APPROVE" and ra:
        rl = [["Facility Amount",f"₹{ra:,.0f}"],["Interest Rate",rate],["Score",f"{ov}/100"],
              ["Type","Working Capital / Term Loan"],["Tenor","12 Months (Renewable)"],
              ["Security","Hypothecation + Property"],["Conditions","Legal docs, CC approval"]]
        rlt = Table(rl, colWidths=[2.2*inch,3.5*inch])
        rlt.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#e4f5ec")),("FONTNAME",(0,0),(0,-1),"Helvetica-Bold"),
            ("GRID",(0,0),(-1,-1),0.5,colors.lightgrey),("PADDING",(0,0),(-1,-1),6),("FONTSIZE",(0,0),(-1,-1),9),
        ]))
        el.append(Spacer(1,6)); el.append(rlt)

    # Annexure
    el.append(PageBreak())
    el.append(Paragraph("ANNEXURE — SCORING METHODOLOGY", h1_s))
    md = [["Category","Weight","Assessment Factors"],
          ["Character","25%","Promoter, NCLT/IBC, MCA, research risk"],
          ["Capacity","30%","DSCR, revenue, EBITDA, GST fraud, circular trading"],
          ["Capital","20%","D/E, net worth, leverage"],
          ["Collateral","15%","Security coverage ratio"],
          ["Conditions","10%","Sector, RBI regulations, macro"]]
    mdt = Table(md, colWidths=[1.3*inch,.8*inch,3.8*inch])
    mdt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),navy),("TEXTCOLOR",(0,0),(-1,0),colors.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("GRID",(0,0),(-1,-1),0.5,colors.lightgrey),("PADDING",(0,0),(-1,-1),5),("FONTSIZE",(0,0),(-1,-1),9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white,colors.HexColor("#f8f8f8")]),
    ]))
    el.append(mdt); el.append(Spacer(1,8))
    el.append(Paragraph("APPROVE (≥70) | REFER (50–70) | REJECT (&lt;50 or hard-reject)", body_s))
    el.append(Paragraph("Hard Rejects: NCLT/IBC, Circular trading, Character &lt;30, Capacity &lt;25", body_s))
    el.append(Spacer(1,20))
    el.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey))
    el.append(Paragraph(f"Generated by Intelli-Credit AI Engine | {dt} | CONFIDENTIAL", small_s))

    pdf.build(el)
    return output_path


def _rating_badge(score: float) -> str:
    if score >= 75: return "STRONG"
    elif score >= 50: return "MODERATE"
    else: return "WEAK"
